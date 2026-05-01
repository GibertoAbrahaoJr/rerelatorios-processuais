from pathlib import Path
from datetime import datetime
from docx import Document
import re

REPORT_DIR = Path("app/storage/reports")
REPORT_DIR.mkdir(parents=True, exist_ok=True)

def safe_filename(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]", "_", text)

def generate_docx(data: dict, analysis: dict) -> str:
    doc = Document()
    doc.add_heading("Relatório Processual Padronizado", 0)

    doc.add_paragraph("Relatório gerado automaticamente a partir da planilha enviada e da consulta pública configurada no sistema.")

    doc.add_heading("1. Identificação do Processo", level=1)
    fields = [
        ("Cliente", data.get("cliente")),
        ("Pasta", data.get("pasta")),
        ("Processo", data.get("numero_processo")),
        ("Tribunal", data.get("tribunal")),
        ("Órgão julgador", data.get("orgao_julgador")),
        ("Classe", data.get("classe")),
        ("Assunto", data.get("assunto")),
        ("Grau/Nível", data.get("grau")),
        ("Data de distribuição/ajuizamento", data.get("data_distribuicao")),
        ("Fonte de consulta", data.get("fonte")),
        ("Status da consulta", data.get("status_consulta")),
        ("Última atualização da consulta", data.get("ultima_atualizacao")),
    ]
    for label, value in fields:
        if value:
            doc.add_paragraph(f"{label}: {value}")

    doc.add_heading("2. Resumo Executivo", level=1)
    doc.add_paragraph(analysis.get("resumo_executivo", ""))

    doc.add_heading("3. Situação Atual", level=1)
    doc.add_paragraph(analysis.get("situacao_atual", ""))

    doc.add_heading("4. Últimos Andamentos Localizados", level=1)
    movimentos = data.get("movimentacoes", [])
    if movimentos:
        for mov in movimentos:
            if isinstance(mov, dict):
                text = f"{mov.get('data', '')} - {mov.get('texto', '')}".strip(" -")
            else:
                text = str(mov)
            doc.add_paragraph(text, style="List Bullet")
    else:
        doc.add_paragraph("Nenhuma movimentação retornada pela fonte pública consultada.")

    doc.add_heading("5. Pontos de Atenção", level=1)
    doc.add_paragraph(analysis.get("pontos_de_atencao", ""))

    doc.add_heading("6. Recomendação", level=1)
    doc.add_paragraph(analysis.get("recomendacao", ""))

    if data.get("observacao"):
        doc.add_heading("7. Observações Internas", level=1)
        doc.add_paragraph(data.get("observacao"))

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    name = safe_filename(data.get("numero_processo", "processo"))
    path = REPORT_DIR / f"relatorio_{name}_{stamp}.docx"
    doc.save(path)
    return str(path)
